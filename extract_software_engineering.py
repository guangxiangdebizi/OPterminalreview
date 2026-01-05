#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
软件工程复习资料提取脚本
从docx文件中提取文字和图片内容
"""

import os
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import io
from PIL import Image

def extract_images_from_docx(docx_path, output_folder):
    """从docx文件中提取所有图片"""
    doc = Document(docx_path)
    image_count = 0
    images_info = []
    
    # 创建图片输出文件夹
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    
    # 遍历文档中的所有关系，查找图片
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_count += 1
            image_data = rel.target_part.blob
            
            # 确定图片扩展名
            ext = rel.target_ref.split('.')[-1]
            if ext not in ['jpg', 'jpeg', 'png', 'gif', 'bmp']:
                ext = 'png'
            
            # 保存图片
            image_filename = f"image_{image_count}.{ext}"
            image_path = os.path.join(output_folder, image_filename)
            
            with open(image_path, 'wb') as f:
                f.write(image_data)
            
            images_info.append({
                'filename': image_filename,
                'path': image_path
            })
            print(f"提取图片: {image_filename}")
    
    return images_info

def extract_text_and_structure(docx_path):
    """提取docx文件的文字内容和结构"""
    doc = Document(docx_path)
    content = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # 判断样式
            style = para.style.name
            content.append({
                'type': 'paragraph',
                'text': text,
                'style': style
            })
    
    # 提取表格
    for table_idx, table in enumerate(doc.tables):
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            table_data.append(row_data)
        content.append({
            'type': 'table',
            'data': table_data
        })
    
    return content

def content_to_markdown(content, images_info, doc_name):
    """将提取的内容转换为Markdown格式"""
    markdown_lines = [f"# {doc_name}\n"]
    image_idx = 0
    
    for item in content:
        if item['type'] == 'paragraph':
            text = item['text']
            style = item['style']
            
            # 根据样式判断标题层级
            if 'Heading 1' in style or '标题 1' in style:
                markdown_lines.append(f"\n## {text}\n")
            elif 'Heading 2' in style or '标题 2' in style:
                markdown_lines.append(f"\n### {text}\n")
            elif 'Heading 3' in style or '标题 3' in style:
                markdown_lines.append(f"\n#### {text}\n")
            else:
                markdown_lines.append(f"{text}\n")
        
        elif item['type'] == 'table':
            markdown_lines.append("\n")
            for row_idx, row in enumerate(item['data']):
                markdown_lines.append("| " + " | ".join(row) + " |")
                if row_idx == 0:
                    markdown_lines.append("| " + " | ".join(["---"] * len(row)) + " |")
            markdown_lines.append("\n")
    
    return "\n".join(markdown_lines)

def process_software_engineering_docs():
    """处理软件工程相关文档"""
    base_dir = r"C:\Users\26214\Desktop\OPreview"
    se_dir = os.path.join(base_dir, "软件工程")
    output_dir = os.path.join(base_dir, "软件工程_提取内容")
    
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    # 处理的文档列表
    docs_to_process = [
        "复习参考.docx",
        "UML实践----用例图、顺序图、状态图、类图、.doc"
    ]
    
    all_markdown = []
    
    for doc_name in docs_to_process:
        doc_path = os.path.join(se_dir, doc_name)
        
        if not os.path.exists(doc_path):
            print(f"文件不存在: {doc_path}")
            continue
        
        print(f"\n处理文档: {doc_name}")
        print("=" * 60)
        
        # 为每个文档创建图片文件夹
        doc_basename = os.path.splitext(doc_name)[0]
        image_folder = os.path.join(output_dir, f"{doc_basename}_图片")
        
        try:
            # 提取图片
            images_info = extract_images_from_docx(doc_path, image_folder)
            print(f"提取了 {len(images_info)} 张图片")
            
            # 提取文字和结构
            content = extract_text_and_structure(doc_path)
            print(f"提取了 {len(content)} 个内容块")
            
            # 转换为Markdown
            markdown_content = content_to_markdown(content, images_info, doc_basename)
            all_markdown.append(markdown_content)
            
            # 保存单独的Markdown文件
            md_path = os.path.join(output_dir, f"{doc_basename}.md")
            with open(md_path, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
            print(f"保存Markdown文件: {md_path}")
            
        except Exception as e:
            print(f"处理文档时出错: {e}")
            import traceback
            traceback.print_exc()
    
    return all_markdown

if __name__ == "__main__":
    print("开始提取软件工程文档内容...")
    print("=" * 60)
    
    try:
        all_content = process_software_engineering_docs()
        print("\n" + "=" * 60)
        print("提取完成！")
        print(f"共处理了 {len(all_content)} 个文档")
    except Exception as e:
        print(f"发生错误: {e}")
        import traceback
        traceback.print_exc()

